import subprocess
from pathlib import Path

import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from vibe_viewer.viewers.base import ViewerError
from vibe_viewer.viewers.document import DocumentViewer
from vibe_viewer.viewers.image import ImageViewer
from vibe_viewer.viewers.office import OfficeDocumentViewer
from vibe_viewer.viewers.spreadsheet import SpreadsheetViewer


def test_png_image_is_rendered(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "image.png"
    Image.new("RGB", (32, 24), "navy").save(path)
    viewer = ImageViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert viewer._pixmap is not None
    assert viewer._pixmap.size().width() == 32


def test_disguised_eps_never_starts_ghostscript(qtbot, monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "disguised.png"
    path.write_bytes(
        b"%!PS-Adobe-3.0 EPSF-3.0\n"
        b"%%BoundingBox: 0 0 10 10\n"
        b"newpath 0 0 moveto 10 10 lineto stroke\nshowpage\n%%EOF\n"
    )
    external_calls = []

    def forbid_external_process(*args, **kwargs):
        external_calls.append((args, kwargs))
        raise AssertionError("An external program was requested")

    monkeypatch.setattr(subprocess, "check_call", forbid_external_process)
    viewer = ImageViewer()
    qtbot.addWidget(viewer)
    with pytest.raises(ViewerError):
        viewer.load_file(path)
    assert external_calls == []


def test_pdf_page_is_rendered(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "document.pdf"
    document = pymupdf.open()
    document.new_page().insert_text((40, 60), "PDF preview")
    document.save(path)
    document.close()
    viewer = DocumentViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert viewer.page_number.maximum() == 1
    assert viewer.page_label.pixmap() is not None
    viewer.unload()


def test_postscript_has_builtin_structural_preview(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "document.ps"
    path.write_text(
        "%!PS-Adobe-3.0\n"
        "%%Title: Vibe demo\n"
        "%%Pages: 1\n"
        "%%BoundingBox: 0 0 100 100\n"
        "%%Page: 1 1\n"
        "/Times-Roman findfont 12 scalefont setfont\n"
        "10 50 moveto (Visible PostScript text) show\n"
        "showpage\n",
        encoding="ascii",
    )
    viewer = DocumentViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert "PostScript" in viewer.info.text()
    assert "Vibe demo" in viewer.page_label.text()
    assert "Visible PostScript text" in viewer.page_label.text()


def test_docx_content_is_rendered(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "document.docx"
    document = Document()
    document.add_heading("DOCX preview", level=1)
    document.add_paragraph("Visible document content")
    document.save(path)
    viewer = OfficeDocumentViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert "Visible document content" in viewer.browser.toPlainText()


def test_xlsx_sheet_is_rendered(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "workbook.xlsx"
    workbook = Workbook()
    workbook.active.append(["name", "score"])
    workbook.active.append(["PDF", 10])
    workbook.save(path)
    viewer = SpreadsheetViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert viewer.table.rowCount() == 2
    assert viewer.table.item(1, 0).text() == "PDF"
    viewer.unload()


def test_dif_table_is_decoded(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "table.dif"
    path.write_text(
        'TABLE\n0,1\n"Vibe"\nVECTORS\n0,2\n""\nTUPLES\n0,2\n""\n'
        'DATA\n0,0\n""\n-1,0\nBOT\n1,0\n"Name"\n1,0\n"Score"\n'
        '-1,0\nBOT\n1,0\n"PDF"\n0,10\nV\n-1,0\nEOD\n',
        encoding="ascii",
    )
    viewer = SpreadsheetViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert viewer.table.rowCount() == 2
    assert viewer.table.item(0, 0).text() == "Name"
    assert viewer.table.item(1, 1).text() == "10"


def test_sylk_table_is_decoded(qtbot, tmp_path: Path) -> None:
    path = tmp_path / "table.slk"
    path.write_text(
        'ID;P\nC;X1;Y1;K"Name"\nC;X2;Y1;K"Score"\n'
        'C;X1;Y2;K"PNG"\nC;X2;Y2;K9\nE\n',
        encoding="ascii",
    )
    viewer = SpreadsheetViewer()
    qtbot.addWidget(viewer)
    viewer.load_file(path)
    assert viewer.table.rowCount() == 2
    assert viewer.table.item(0, 0).text() == "Name"
    assert viewer.table.item(1, 1).text() == "9"
